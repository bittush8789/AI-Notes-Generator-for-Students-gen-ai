import os
import re
import zipfile
import xml.etree.ElementTree as ET
import logging

# Check for PDF parsing libraries
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import pypdf
except ImportError:
    try:
        import PyPDF2 as pypdf
    except ImportError:
        pypdf = None

# Check for Word parsing
try:
    import docx
except ImportError:
    docx = None

logger = logging.getLogger(__name__)

def clean_text(text: str) -> str:
    """Cleans up excessive whitespaces, newlines, and non-printable characters."""
    if not text:
        return ""
    # Normalize whitespaces
    text = re.sub(r'[ \t]+', ' ', text)
    # Normalize consecutive empty lines (allow max 2 newlines)
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()

def parse_txt(filepath: str) -> str:
    """Parses a standard plain text file."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        logger.error(f"Failed to parse text file: {e}")
        raise ValueError(f"Failed to parse TXT file: {str(e)}")

def parse_docx(filepath: str) -> str:
    """Parses DOCX documents using python-docx."""
    if docx is None:
        raise ValueError("python-docx is not installed.")
    
    try:
        doc = docx.Document(filepath)
        full_text = []
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Failed to parse DOCX: {e}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")

def parse_pptx(filepath: str) -> str:
    """
    Parses a PPTX presentation file using the standard zip/XML parsing.
    This is extremely robust and does not require third-party libraries.
    """
    try:
        texts = []
        with zipfile.ZipFile(filepath, 'r') as zip_ref:
            # Sort files in ppt/slides/
            slide_files = sorted(
                [f for f in zip_ref.namelist() if re.match(r'ppt/slides/slide\d+\.xml', f)],
                key=lambda x: int(re.search(r'\d+', x).group())
            )
            
            # Namespace for pptx elements
            namespace = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}
            
            for index, slide_file in enumerate(slide_files, start=1):
                texts.append(f"\n--- Slide {index} ---")
                slide_xml = zip_ref.read(slide_file)
                root = ET.fromstring(slide_xml)
                
                # Find all text elements in the slide XML
                # In pptx, texts are inside <a:t> elements
                slide_texts = []
                for text_elem in root.findall('.//a:t', namespaces=namespace):
                    if text_elem.text:
                        slide_texts.append(text_elem.text)
                
                texts.append("\n".join(slide_texts))
                
        return "\n".join(texts)
    except Exception as e:
        logger.error(f"Failed to parse PPTX file: {e}")
        raise ValueError(f"Failed to parse PPTX: {str(e)}")

def parse_pdf(filepath: str) -> str:
    """
    Parses PDF using pdfplumber as priority (handles tables and structure well)
    and pypdf as fallback.
    """
    pdf_text = []

    # Method 1: pdfplumber
    if pdfplumber is not None:
        try:
            logger.info("Parsing PDF with pdfplumber...")
            with pdfplumber.open(filepath) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        pdf_text.append(text)
            if pdf_text:
                return "\n\n".join(pdf_text)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}. Retrying with pypdf...")

    # Method 2: pypdf fallback
    if pypdf is not None:
        try:
            logger.info("Parsing PDF with pypdf...")
            with open(filepath, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        pdf_text.append(text)
            if pdf_text:
                return "\n\n".join(pdf_text)
        except Exception as e:
            logger.error(f"pypdf fallback failed: {e}")
            raise ValueError(f"Failed to parse PDF using available libraries: {str(e)}")

    raise ValueError("No PDF parsing libraries (pdfplumber, pypdf) are available/successful.")

def parse_file(filepath: str, file_type: str) -> str:
    """Main router to parse files and return cleaned text content."""
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")
    
    file_type = file_type.lower()
    
    if file_type == "txt":
        raw_text = parse_txt(filepath)
    elif file_type == "docx":
        raw_text = parse_docx(filepath)
    elif file_type == "pptx" or file_type == "ppt":
        raw_text = parse_pptx(filepath)
    elif file_type == "pdf":
        raw_text = parse_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file format: {file_type}")
    
    cleaned = clean_text(raw_text)
    if not cleaned:
        raise ValueError("Could not extract any readable text from the file.")
    
    return cleaned
